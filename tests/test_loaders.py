"""loaders.py 단위 테스트.

**픽스처 정책이 한 번 바뀌었다(2026-08).** 원래 이 파일은 "파서는 실물
문서로만 검증한다"는 원칙에 따라 실제 이력서 DOCX와 국토교통부 보도자료
HWPX를 `tests/fixtures/`에 두고 썼다. 그런데 그 둘은 각각 **개인정보와
외부 기관 문서**라 공개 저장소에 있으면 안 되는 파일이었고, 히스토리
정리로 제거됐다(`.gitignore` 참고).

문제는 그 다음이었다 — 파일이 사라졌는데 테스트는 `skipif`로 **조용히
건너뛰었다**. CI는 초록불인데 DOCX·HWPX 파서는 아무도 검증하지 않는
상태가 됐다(실측: 121 passed → 119 passed + 2 skipped). "검증했다"고
문서에 적힌 채 검증이 사라지는 게 가장 나쁜 조합이라, 픽스처를 **테스트
시점에 코드로 생성**하도록 바꿨다.

생성물은 가짜가 아니다 — `python-docx`·`python-hwpx`가 각 규격대로 쓴
진짜 OOXML/HWPX 파일이다(HWPX는 ZIP 매직바이트 `PK\\x03\\x04`까지
확인한다). 다만 **원래 검증보다 약하다**는 점은 정직하게 남긴다:

  - 같은 라이브러리로 쓰고 읽는 왕복(round-trip)이라, 한글(HWP) 같은
    **다른 생산자가 만든 파일**의 변형은 못 잡는다. 실제로 이 프로젝트는
    "HWPX인 줄 알았는데 OLE 기반 구형 .hwp였던" 파일을 두 번 만났고,
    그건 실물 검증이었기에 잡혔다.
  - 그 커버리지는 이제 없다. 대신 매직바이트 검사로 최소한 "ZIP 기반
    HWPX가 맞는지"는 계속 확인한다.

PDF만은 실물을 그대로 쓴다 — 공개된 학술 논문이라 저장소에 둬도 되고,
실제 코퍼스(`data/docs/`)에 들어가 있는 파일이기도 하다.
"""
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

from loaders import LOADERS, _normalize_whitespace, load_docx, load_hwpx, load_pdf  # noqa: E402

DOCS_DIR = ROOT / "data" / "docs"
PDF_PATH = DOCS_DIR / "segmentation-paper.pdf"


@pytest.fixture(scope="module")
def docx_path(tmp_path_factory) -> Path:
    """python-docx로 진짜 .docx를 만든다(문단 + 표) — 로더가 둘 다 읽는지 본다."""
    from docx import Document as DocxDocument

    path = tmp_path_factory.mktemp("fixtures") / "sample-resume.docx"
    d = DocxDocument()
    d.add_paragraph("이윤선 — AI 엔지니어")
    d.add_paragraph("TTS 프로젝트: TTFB 2292ms → 334ms 개선")
    # 표도 넣는다. load_docx가 문단뿐 아니라 표 행까지 훑는지 검증하는
    # 유일한 지점이라, 표를 빼면 그 분기가 통째로 미검증이 된다.
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "회사"
    table.cell(0, 1).text = "기간"
    table.cell(1, 0).text = "MiCo AI"
    table.cell(1, 1).text = "2025.04~"
    d.save(path)
    return path


@pytest.fixture(scope="module")
def hwpx_path(tmp_path_factory) -> Path:
    """python-hwpx로 규격에 맞는 진짜 .hwpx(ZIP+XML)를 만든다."""
    from hwpx.document import HwpxDocument

    path = tmp_path_factory.mktemp("fixtures") / "sample-notice.hwpx"
    doc = HwpxDocument.new()
    doc.add_paragraph("국토교통부 보도자료")
    doc.add_paragraph("제3회 철도의 날 기념식을 개최한다.")
    doc.save_to_path(path)
    return path


def test_normalize_whitespace_collapses_and_trims():
    assert _normalize_whitespace("a  \nb\n\n\n\nc  \n") == "a\nb\n\nc"


@pytest.mark.skipif(not PDF_PATH.exists(), reason="실물 PDF 없음")
def test_load_pdf_extracts_real_content():
    doc = load_pdf(PDF_PATH)
    assert doc.metadata["source"] == "segmentation-paper.pdf"
    assert doc.metadata["pages"] > 0
    assert len(doc.page_content) > 500
    # 논문 핵심 키워드가 실제로 뽑혔는지 — 빈 페이지 껍데기만 남는 회귀 방지
    assert "세그멘테이션" in doc.page_content


def test_load_docx_extracts_paragraphs_and_table_cells(docx_path):
    doc = load_docx(docx_path)
    assert doc.metadata["source"] == "sample-resume.docx"
    assert "이윤선" in doc.page_content
    assert "2292" in doc.page_content
    # 표 셀까지 텍스트로 들어왔는지 — load_docx의 표 분기 검증
    assert "MiCo AI" in doc.page_content
    assert "2025.04~" in doc.page_content


def test_load_hwpx_reads_real_zip_based_package(hwpx_path):
    # 이 프로젝트는 "HWPX인 줄 알았는데 OLE 구형 .hwp"인 파일을 두 번 만났다.
    # 생성물이 ZIP 기반 HWPX가 맞는지부터 고정한다.
    assert hwpx_path.read_bytes()[:4] == b"PK\x03\x04"

    doc = load_hwpx(hwpx_path)
    assert doc.metadata["source"] == "sample-notice.hwpx"
    assert "국토교통부" in doc.page_content
    assert "철도의 날" in doc.page_content


def test_loaders_registry_covers_pdf_docx_hwpx():
    assert set(LOADERS) == {".pdf", ".docx", ".hwpx"}
