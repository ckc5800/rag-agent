"""비마크다운 포맷 로더 — PDF·DOCX·HWPX. ingest.py의 load_documents()가 이
함수들을 불러 마크다운과 같은 파이프라인(정제→청킹→임베딩)에 합류시킨다.

채용 공고(그래파이 등)가 요구한 "비정형 데이터(PDF/DOCX/HWP) 처리"의 실제
구현. 구형 바이너리 .hwp는 만들지 않았다 — 이 PC·계정 전체를 검색해도 실물
파일이 0건이고(2026-08 확인), 검증할 실물 없이 파서를 만들면 동작을 아무도
확인 못 한다. 대신 **신형 .hwpx**(ZIP+XML 기반, 2026-05-18부터 중앙부처·
지자체 공문서 표준으로 의무화된 포맷)는 지원한다 — 개인 소유 실물은 없지만,
국토교통부가 실제로 배포한 보도자료(2026 철도의 날 기념행사, 공개 웹에서
직접 내려받음)로 검증했다. 합성 문서가 아니라는 이 프로젝트의 원칙은
"내가 쓴 문서"가 아니라 "실재하는 문서"로 지켜도 된다 — KLUE-RE 실험에서
개인 코퍼스가 작아 공개 벤치마크로 검증 범위를 넓혔던 것과 같은 논리다.
"""
import re
from pathlib import Path

from langchain_core.documents import Document

_EXTRA_BLANK = re.compile(r"\n{3,}")
_TRAILING_WS = re.compile(r"[ \t]+\n")


def _normalize_whitespace(text: str) -> str:
    """줄 끝 공백 제거 + 빈 줄 3개 이상을 2개로. PDF 텍스트 추출은 페이지마다
    들쭉날쭉한 공백을 남기는데, 마크다운의 clean_markdown과 달리 이 포맷들엔
    노션 매크로·HTML 태그가 없어 정제할 게 이것뿐이다."""
    text = _TRAILING_WS.sub("\n", text)
    return _EXTRA_BLANK.sub("\n\n", text).strip()


def load_pdf(path: Path) -> Document:
    """PDF 텍스트를 페이지 순서대로 이어붙인다.

    pypdf는 레이아웃을 모르는 순수 텍스트 추출기라 표·다이어그램은 문장이
    아니라 단어가 흩어진 형태로 나온다(발표자료 PDF의 아키텍처 그림 페이지가
    특히 그렇다) — 마크다운의 박스 그림처럼 문자로 판별할 신호(┌─│ 등)가
    없어 자동 분리는 하지 않는다. ingest 후 inspect_data.py로 실제 어떤
    노이즈가 남는지 보고 판단한다(이 프로젝트가 항상 해 온 순서).
    """
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [p.extract_text() or "" for p in reader.pages]
    text = _normalize_whitespace("\n\n".join(pages))
    return Document(page_content=text,
                    metadata={"source": path.name, "pages": len(pages)})


def load_docx(path: Path) -> Document:
    """DOCX 문단을 순서대로 이어붙이고, 표가 있으면 행마다 " | "로 이어
    한 줄씩 덧붙인다 — 문단만 읽으면 표 안 내용이 통째로 빠진다."""
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = _normalize_whitespace("\n".join(parts))
    return Document(page_content=text, metadata={"source": path.name})


def load_hwpx(path: Path) -> Document:
    """HWPX(ZIP+XML) 본문을 평문으로 추출한다. `python-hwpx`(Apache-2.0)에
    위임 — HWPX는 표 셀·문단 구조가 DOCX보다 복잡해(섹션·트랙체인지 등)
    직접 XML을 파싱하는 것보다 검증된 라이브러리를 쓰는 게 안전하다.
    `.text.plain()`을 쓴다(구버전 `export_text()`는 6.0에서 폐기 예고)."""
    import hwpx

    doc = hwpx.HwpxDocument.open(str(path))
    try:
        text = _normalize_whitespace(doc.text.plain())
    finally:
        doc.close()
    return Document(page_content=text, metadata={"source": path.name})


# 확장자 → 로더. ingest.py의 load_documents()가 이 표로 포맷을 분기한다.
LOADERS = {".pdf": load_pdf, ".docx": load_docx, ".hwpx": load_hwpx}
